# utils.py - Utility Functions and Helpers
import os
import re
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import hashlib
import tempfile
from pathlib import Path

import streamlit as st
import PyPDF2
from docx import Document
import nltk
from collections import Counter

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileHandler:
    """Utility class for handling file operations"""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    def validate_file(cls, uploaded_file) -> Tuple[bool, str]:
        """Validate uploaded file"""
        if not uploaded_file:
            return False, "No file provided"
        
        # Check file size
        if uploaded_file.size > cls.MAX_FILE_SIZE:
            return False, f"File too large. Maximum size: {cls.MAX_FILE_SIZE // 1024 // 1024}MB"
        
        # Check file extension
        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in cls.SUPPORTED_FORMATS:
            return False, f"Unsupported format. Supported: {', '.join(cls.SUPPORTED_FORMATS)}"
        
        return True, "File is valid"
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        """Extract text from PDF file with error handling"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            if not text.strip():
                return "Error: Could not extract text from PDF. File may be image-based or corrupted."
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return f"Error reading PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        """Extract text from DOCX file with error handling"""
        try:
            doc = Document(file)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                return "Error: Could not extract text from DOCX file."
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return f"Error reading DOCX: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        """Extract text from TXT file with encoding handling"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'ascii', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.seek(0)  # Reset file pointer
                    content = file.read()
                    if isinstance(content, bytes):
                        text = content.decode(encoding)
                    else:
                        text = str(content)
                    
                    if text.strip():
                        return text.strip()
                        
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Error with encoding {encoding}: {e}")
                    continue
            
            return "Error: Could not decode text file with any supported encoding."
            
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return f"Error reading TXT: {str(e)}"

class TextProcessor:
    """Utility class for text processing and analysis"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep essential punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)\@\#\%\&\+\=]', ' ', text)
        
        # Clean up multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def extract_email(text: str) -> Optional[str]:
        """Extract email address from text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else None
    
    @staticmethod
    def extract_phone(text: str) -> Optional[str]:
        """Extract phone number from text"""
        # Various phone number patterns
        patterns = [
            r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'([0-9]{3})[-.]?([0-9]{3})[-.]?([0-9]{4})'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                # Format the phone number
                match = matches[0]
                if isinstance(match, tuple):
                    return f"({match[0]}) {match[1]}-{match[2]}"
                else:
                    return match
        
        return None
    
    @staticmethod
    def extract_urls(text: str) -> List[str]:
        """Extract URLs from text"""
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        
        urls = re.findall(url_pattern, text)
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        
        all_urls = urls + [f"https://{link}" for link in linkedin]
        return list(set(all_urls))
    
    @staticmethod
    def count_words(text: str) -> Dict[str, int]:
        """Count words and provide text statistics"""
        if not text:
            return {"words": 0, "characters": 0, "lines": 0, "paragraphs": 0}
        
        words = len(text.split())
        characters = len(text)
        lines = len(text.split('\n'))
        paragraphs = len([p for p in text.split('\n\n') if p.strip()])
        
        return {
            "words": words,
            "characters": characters,
            "lines": lines,
            "paragraphs": paragraphs
        }
    
    @staticmethod
    def extract_years(text: str) -> List[int]:
        """Extract years from text (useful for experience dates)"""
        year_pattern = r'\b(19|20)\d{2}\b'
        years = [int(year) for year in re.findall(year_pattern, text)]
        return sorted(list(set(years)))

class ResumeValidator:
    """Validate resume content and structure"""
    
    REQUIRED_SECTIONS = ['contact', 'experience', 'education']
    OPTIONAL_SECTIONS = ['summary', 'skills', 'projects', 'certifications']
    
    @classmethod
    def validate_resume_structure(cls, text: str) -> Dict[str, Any]:
        """Validate basic resume structure"""
        validation_result = {
            "is_valid": True,
            "issues": [],
            "warnings": [],
            "sections_found": [],
            "missing_sections": []
        }
        
        text_lower = text.lower()
        
        # Check for contact information
        if not TextProcessor.extract_email(text):
            validation_result["issues"].append("No email address found")
            validation_result["is_valid"] = False
        
        if not TextProcessor.extract_phone(text):
            validation_result["warnings"].append("No phone number found")
        
        # Check for common resume sections
        section_keywords = {
            'experience': ['experience', 'employment', 'work history', 'professional experience'],
            'education': ['education', 'academic', 'degree', 'university', 'college'],
            'skills': ['skills', 'technical skills', 'competencies', 'technologies'],
            'summary': ['summary', 'profile', 'objective', 'about']
        }
        
        for section, keywords in section_keywords.items():
            found = any(keyword in text_lower for keyword in keywords)
            if found:
                validation_result["sections_found"].append(section)
            elif section in cls.REQUIRED_SECTIONS:
                validation_result["missing_sections"].append(section)
                validation_result["is_valid"] = False
        
        # Check resume length
        word_count = len(text.split())
        if word_count < 100:
            validation_result["issues"].append("Resume appears too short")
            validation_result["is_valid"] = False
        elif word_count > 1000:
            validation_result["warnings"].append("Resume may be too long for single page")
        
        return validation_result

class CacheManager:
    """Manage caching for expensive operations"""
    
    @staticmethod
    def generate_cache_key(content: str, additional_params: Dict = None) -> str:
        """Generate cache key for content"""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        
        if additional_params:
            params_str = json.dumps(additional_params, sort_keys=True)
            params_hash = hashlib.md5(params_str.encode()).hexdigest()
            return f"{content_hash}_{params_hash}"
        
        return content_hash
    
    @staticmethod
    def save_to_cache(key: str, data: Any, cache_dir: str = None) -> bool:
        """Save data to cache"""
        try:
            if cache_dir is None:
                cache_dir = tempfile.gettempdir()
            
            cache_path = Path(cache_dir) / f"resume_cache_{key}.json"
            
            cache_data = {
                "timestamp": datetime.now().isoformat(),
                "data": data
            }
            
            with open(cache_path, 'w') as f:
                json.dump(cache_data, f)
            
            return True
            
        except Exception as e:
            logger.error(f"Cache save error: {e}")
            return False
    
    @staticmethod
    def load_from_cache(key: str, cache_dir: str = None, max_age_hours: int = 24) -> Optional[Any]:
        """Load data from cache"""
        try:
            if cache_dir is None:
                cache_dir = tempfile.gettempdir()
            
            cache_path = Path(cache_dir) / f"resume_cache_{key}.json"
            
            if not cache_path.exists():
                return None
            
            with open(cache_path, 'r') as f:
                cache_data = json.load(f)
            
            # Check if cache is still valid
            cache_time = datetime.fromisoformat(cache_data["timestamp"])
            age_hours = (datetime.now() - cache_time).total_seconds() / 3600
            
            if age_hours > max_age_hours:
                cache_path.unlink()  # Delete expired cache
                return None
            
            return cache_data["data"]
            
        except Exception as e:
            logger.error(f"Cache load error: {e}")
            return None

class StreamlitHelpers:
    """Helper functions for Streamlit UI components"""
    
    @staticmethod
    def display_progress_with_status(steps: List[str], current_step: int = 0):
        """Display progress bar with step descriptions"""
        progress = current_step / len(steps) if steps else 0
        progress_bar = st.progress(progress)
        
        status_container = st.container()
        
        with status_container:
            for i, step in enumerate(steps):
                if i < current_step:
                    st.success(f"✅ {step}")
                elif i == current_step:
                    st.info(f"🔄 {step}")
                else:
                    st.empty()
        
        return progress_bar, status_container
    
    @staticmethod
    def display_metrics_grid(metrics: Dict[str, Any], columns: int = 3):
        """Display metrics in a grid layout"""
        cols = st.columns(columns)
        
        for i, (key, value) in enumerate(metrics.items()):
            col_idx = i % columns
            
            with cols[col_idx]:
                if isinstance(value, dict) and "value" in value:
                    st.metric(
                        key.replace("_", " ").title(),
                        value["value"],
                        value.get("delta", None)
                    )
                else:
                    st.metric(key.replace("_", " ").title(), value)
    
    @staticmethod
    def display_agent_status(agents: Dict[str, str]):
        """Display agent status indicators"""
        st.subheader("🤖 Agent Status")
        
        for agent_name, status in agents.items():
            if status.lower() == "ready":
                st.success(f"✅ {agent_name.replace('_', ' ').title()}: {status}")
            elif status.lower() == "processing":
                st.info(f"🔄 {agent_name.replace('_', ' ').title()}: {status}")
            elif status.lower() == "error":
                st.error(f"❌ {agent_name.replace('_', ' ').title()}: {status}")
            else:
                st.warning(f"⚠️ {agent_name.replace('_', ' ').title()}: {status}")
    
    @staticmethod
    def create_download_button(content: str, filename: str, label: str = "Download"):
        """Create a download button for content"""
        return st.download_button(
            label=f"📥 {label}",
            data=content,
            file_name=filename,
            mime="text/plain"
        )

class ErrorHandler:
    """Centralized error handling"""
    
    @staticmethod
    def handle_api_error(error: Exception, context: str = "API call") -> str:
        """Handle API-related errors"""
        error_msg = str(error)
        
        if "api key" in error_msg.lower():
            return f"❌ API Key Error: Please check your API key configuration"
        elif "rate limit" in error_msg.lower():
            return f"⏳ Rate Limit: Please wait before making another request"
        elif "timeout" in error_msg.lower():
            return f"⏱️ Timeout: The request took too long. Please try again"
        else:
            logger.error(f"{context} error: {error}")
            return f"❌ {context} Error: {error_msg}"
    
    @staticmethod
    def handle_file_error(error: Exception, filename: str) -> str:
        """Handle file processing errors"""
        error_msg = str(error)
        
        if "permission" in error_msg.lower():
            return f"🔒 Permission Error: Cannot access file {filename}"
        elif "not found" in error_msg.lower():
            return f"📁 File Not Found: {filename}"
        elif "corrupt" in error_msg.lower():
            return f"💥 Corrupt File: {filename} appears to be damaged"
        else:
            logger.error(f"File processing error for {filename}: {error}")
            return f"❌ File Error: Could not process {filename} - {error_msg}"
    
    @staticmethod
    def display_error_with_help(error_msg: str, help_text: str = None):
        """Display error with helpful suggestions"""
        st.error(error_msg)
        
        if help_text:
            st.info(f"💡 {help_text}")
        else:
            # Provide generic help based on error type
            if "api key" in error_msg.lower():
                st.info("💡 Make sure to enter your OpenAI API key in the sidebar")
            elif "file" in error_msg.lower():
                st.info("💡 Try uploading a different file format (PDF, DOCX, or TXT)")
            elif "timeout" in error_msg.lower():
                st.info("💡 Try reducing the file size or check your internet connection")

# Utility functions for data processing
def format_file_size(size_bytes: int) -> str:
    """Format file size in human readable format"""
    if size_bytes < 1024:
        return f"{size_bytes} bytes"
    elif size_bytes < 1024**2:
        return f"{size_bytes/1024:.1f} KB"
    else:
        return f"{size_bytes/(1024**2):.1f} MB"

def estimate_reading_time(text: str) -> int:
    """Estimate reading time in minutes"""
    word_count = len(text.split())
    # Average reading speed: 200-250 words per minute
    return max(1, word_count // 225)

def generate_resume_filename(original_name: str, suffix: str = "optimized") -> str:
    """Generate filename for optimized resume"""
    name_part = Path(original_name).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{name_part}_{suffix}_{timestamp}.txt"